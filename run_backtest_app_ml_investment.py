# ============================================================
# RUN BACKTEST - APP ML INVESTMENT
# 2016 → 2025
# ============================================================

import logging
from pathlib import Path
import numpy as np
import pandas as pd

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from database.connection import DatabaseConnection
from preprocessing.data_preprocessor import DataPreprocessor
from models.randomForest import RandomForestModel


# =========================
# CONFIGURAÇÃO FIXA
# =========================

START_YEAR = 2016
END_YEAR = 2025
TOP_N = 20

DB_HOST = "localhost"
DB_USER = "root"
DB_PASSWORD = "Oracle1991#"
DB_NAME = "finance_db"
DB_PORT = 3306


# ============================================================
# LOAD DATA
# ============================================================

def load_data(conn):
    query = """
    SELECT
        t.ticker,
        f.ano,
        f.dy, f.pl, f.peg_ratio, f.pvp, f.ev_ebitda, f.ev_ebit,
        f.p_ebitda, f.p_ebit, f.vpa, f.p_ativo, f.lpa, f.psr,
        f.p_cap_giro, f.p_ativo_circ_liq, f.div_liquida_pl,
        f.div_liquida_ebitda, f.div_liquida_ebit, f.pl_ativos,
        f.passivos_ativos, f.liq_corrente, f.m_bruta, f.m_ebitda,
        f.m_ebit, f.m_liquida, f.roe, f.roa, f.roic, f.giro_ativos,
        f.cagr_receitas_5y, f.cagr_lucros_5y,
        r.retorno
    FROM fundamentals f
    JOIN dim_tickers t ON f.ticker_id = t.id
    LEFT JOIN returns r
        ON r.ticker_id = f.ticker_id
       AND r.ano_inicio = f.ano
       AND r.ano_fim = f.ano + 1
    WHERE f.ano BETWEEN %s AND %s
    ORDER BY t.ticker, f.ano
    """

    cur = conn.cursor(dictionary=True)
    cur.execute(query, (START_YEAR, END_YEAR))
    rows = cur.fetchall()
    cur.close()

    df = pd.DataFrame(rows)

    # Converter DECIMAL para float
    for col in df.columns:
        if col not in ["ticker", "ano"]:
            df[col] = pd.to_numeric(df[col], errors="coerce")

    return df


# ============================================================
# WALK FORWARD
# ============================================================

# ============================================================
# RUN BACKTEST - APP ML INVESTMENT
# 2016 → 2025
# ============================================================

import logging
from pathlib import Path
import numpy as np
import pandas as pd

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from database.connection import DatabaseConnection
from preprocessing.data_preprocessor import DataPreprocessor
from models.randomForest import RandomForestModel


# =========================
# CONFIGURAÇÃO FIXA
# =========================

START_YEAR = 2016
END_YEAR = 2025
TOP_N = 20

DB_HOST = "localhost"
DB_USER = "root"
DB_PASSWORD = "Oracle1991#"
DB_NAME = "finance_db"
DB_PORT = 3306


# ============================================================
# LOAD DATA
# ============================================================

def load_data(conn):
    query = """
    SELECT
        t.ticker,
        f.ano,
        f.dy, f.pl, f.peg_ratio, f.pvp, f.ev_ebitda, f.ev_ebit,
        f.p_ebitda, f.p_ebit, f.vpa, f.p_ativo, f.lpa, f.psr,
        f.p_cap_giro, f.p_ativo_circ_liq, f.div_liquida_pl,
        f.div_liquida_ebitda, f.div_liquida_ebit, f.pl_ativos,
        f.passivos_ativos, f.liq_corrente, f.m_bruta, f.m_ebitda,
        f.m_ebit, f.m_liquida, f.roe, f.roa, f.roic, f.giro_ativos,
        f.cagr_receitas_5y, f.cagr_lucros_5y,
        r.retorno
    FROM fundamentals f
    JOIN dim_tickers t ON f.ticker_id = t.id
    LEFT JOIN returns r
        ON r.ticker_id = f.ticker_id
       AND r.ano_inicio = f.ano
       AND r.ano_fim = f.ano + 1
    WHERE f.ano BETWEEN %s AND %s
    ORDER BY t.ticker, f.ano
    """

    cur = conn.cursor(dictionary=True)
    cur.execute(query, (START_YEAR, END_YEAR))
    rows = cur.fetchall()
    cur.close()

    df = pd.DataFrame(rows)

    # Converter DECIMAL para float
    for col in df.columns:
        if col not in ["ticker", "ano"]:
            df[col] = pd.to_numeric(df[col], errors="coerce")

    return df


# ============================================================
# WALK FORWARD
# ============================================================

def run_backtest(df):

    df = df.sort_values(["ticker", "ano"]).copy()

    # Criar target = retorno do ano seguinte
    df["target"] = df.groupby("ticker")["retorno"].shift(-1)
    df = df.dropna(subset=["target"]).reset_index(drop=True)

    pre = DataPreprocessor(year_col="ano", ticker_col="ticker")
    df = pre.impute_median_by_year(df)

    feature_cols = [
        col for col in df.columns
        if col not in ["ticker", "ano", "retorno", "target"]
    ]

    rf = RandomForestModel(
        n_estimators=500,
        min_samples_leaf=2,
        random_state=42,
        n_jobs=-1
    )

    results = []
    details = {}

    for year in range(START_YEAR, END_YEAR):

        train_df = df[df["ano"] == year]
        test_df = df[df["ano"] == year + 1]

        if train_df.empty or test_df.empty:
            continue

        X_train = train_df[feature_cols]
        y_train = train_df["target"]

        X_test = test_df[feature_cols]
        y_test = test_df["target"]

        rf.train(X_train, y_train)
        preds = rf.predict(X_test)

        test_df = test_df.copy()
        test_df["prediction"] = preds

        ranked = test_df.sort_values("prediction", ascending=False)
        top = ranked.head(TOP_N)

        portfolio_return = top["target"].mean()
        benchmark_return = ranked["target"].mean()

        results.append({
            "train_year": year,
            "test_year": year + 1,
            "portfolio_return": portfolio_return,
            "benchmark_return": benchmark_return,
            "alpha": portfolio_return - benchmark_return
        })

        details[year + 1] = ranked

    return pd.DataFrame(results), details


# ============================================================
# EXPORT XLSX
# ============================================================

def export_xlsx(summary_df, details):
    wb = Workbook()
    ws = wb.active
    ws.title = "resumo"

    for r in dataframe_to_rows(summary_df, index=False, header=True):
        ws.append(r)

    for year, df in details.items():
        sheet = wb.create_sheet(f"det_{year}")
        for r in dataframe_to_rows(df, index=False, header=True):
            sheet.append(r)

    wb.save("backtest_random_forest.xlsx")


# ============================================================
# EXPORT DOCX
# ============================================================

def export_docx(summary_df):
    doc = Document()

    title = doc.add_heading("Relatório Backtest Random Forest", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if summary_df.empty:
        doc.add_paragraph("Nenhum resultado disponível.")
        doc.save("backtest_random_forest_report.docx")
        return

    table = doc.add_table(rows=1, cols=len(summary_df.columns))

    for i, col in enumerate(summary_df.columns):
        table.rows[0].cells[i].text = col

    for _, row in summary_df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(summary_df.columns):
            val = row[col]
            if isinstance(val, float):
                cells[i].text = f"{val:.2%}"
            else:
                cells[i].text = str(val)

    doc.save("backtest_random_forest_report.docx")


# ============================================================
# MAIN
# ============================================================

def main():
    logging.basicConfig(level=logging.INFO)

    db = DatabaseConnection(
        host=DB_HOST,
        user=DB_USER,
        password=DB_PASSWORD,
        database=DB_NAME,
        port=DB_PORT
    )

    conn = db.get_connection()

    df = load_data(conn)

    summary, details = run_backtest(df)

    export_xlsx(summary, details)
    export_docx(summary)

    print("\n==== RESULTADO ====\n")
    print(summary)

    db.close()


if __name__ == "__main__":
    main()


# ============================================================
# EXPORT XLSX
# ============================================================

def export_xlsx(summary_df, details):
    wb = Workbook()
    ws = wb.active
    ws.title = "resumo"

    for r in dataframe_to_rows(summary_df, index=False, header=True):
        ws.append(r)

    for year, df in details.items():
        sheet = wb.create_sheet(f"det_{year}")
        for r in dataframe_to_rows(df, index=False, header=True):
            sheet.append(r)

    wb.save("backtest_random_forest.xlsx")


# ============================================================
# EXPORT DOCX
# ============================================================

def export_docx(summary_df):
    doc = Document()

    title = doc.add_heading("Relatório Backtest Random Forest", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if summary_df.empty:
        doc.add_paragraph("Nenhum resultado disponível.")
        doc.save("backtest_random_forest_report.docx")
        return

    table = doc.add_table(rows=1, cols=len(summary_df.columns))

    for i, col in enumerate(summary_df.columns):
        table.rows[0].cells[i].text = col

    for _, row in summary_df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(summary_df.columns):
            val = row[col]
            if isinstance(val, float):
                cells[i].text = f"{val:.2%}"
            else:
                cells[i].text = str(val)

    doc.save("backtest_random_forest_report.docx")


# ============================================================
# MAIN
# ============================================================

def main():
    logging.basicConfig(level=logging.INFO)

    db = DatabaseConnection(
        host=DB_HOST,
        user=DB_USER,
        password=DB_PASSWORD,
        database=DB_NAME,
        port=DB_PORT
    )

    conn = db.get_connection()

    df = load_data(conn)

    summary, details = run_backtest(df)

    export_xlsx(summary, details)
    export_docx(summary)

    print("\n==== RESULTADO ====\n")
    print(summary)

    db.close()


if __name__ == "__main__":
    main()